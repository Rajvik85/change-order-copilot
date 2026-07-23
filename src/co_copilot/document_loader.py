"""Document ingestion and best-effort section segmentation."""

from __future__ import annotations

import hashlib
import io
import logging
import re
import unicodedata
import zipfile
from pathlib import Path

from docx import Document as WordDocument

from co_copilot.exceptions import DocumentParseError
from co_copilot.models import Document

LOGGER = logging.getLogger(__name__)
SUPPORTED_EXTENSIONS = {".md", ".txt", ".docx"}

_SECTION_ALIASES = {
    "header": "header",
    "event description": "description",
    "description": "description",
    "event": "description",
    "contractual basis": "basis",
    "basis": "basis",
    "cost impact": "cost_impact",
    "schedule impact": "schedule_impact",
    "impacts": "impacts",
    "attachments": "attachments",
}


def normalize_text(text: str) -> str:
    """Normalize Unicode and whitespace while preserving line structure."""
    normalized = unicodedata.normalize("NFKC", text).replace("\ufeff", "")
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in normalized.splitlines()]
    return "\n".join(lines).strip()


def _decode_text(data: bytes, filename: str) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError(
        f"{filename} could not be decoded as UTF-8 or Windows-1252 text."
    )


def _read_docx(data: bytes, filename: str) -> str:
    try:
        document = WordDocument(io.BytesIO(data))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    except (ValueError, KeyError, OSError, zipfile.BadZipFile) as exc:
        raise DocumentParseError(
            f"{filename} is not a readable Word document. Re-save it as .docx."
        ) from exc


def segment_sections(text: str) -> tuple[dict[str, str], tuple[str, ...]]:
    """Segment common Markdown or plain-text headings.

    Unknown layouts deliberately degrade to a full-text section so extraction
    can still attempt labeled-field and global regex patterns.
    """
    sections: dict[str, list[str]] = {}
    active: str | None = None
    for line in text.splitlines():
        candidate = re.sub(r"^[#>*\-\s]+", "", line).strip().rstrip(":")
        key = _SECTION_ALIASES.get(candidate.lower())
        if key or (
            candidate.isupper()
            and (key := _SECTION_ALIASES.get(candidate.lower())) is not None
        ):
            active = key
            sections.setdefault(active, [])
            continue
        if active is not None:
            sections[active].append(line)

    compact = {
        name: "\n".join(lines).strip()
        for name, lines in sections.items()
        if "\n".join(lines).strip()
    }
    warnings: list[str] = []
    if not compact:
        compact = {"full_text": text}
        warning = (
            "No recognized section headings; analysis continued in full-text mode."
        )
        warnings.append(warning)
        LOGGER.warning(warning)
    elif "cost_impact" in compact or "schedule_impact" in compact:
        impacts = "\n".join(
            value
            for key, value in compact.items()
            if key in {"cost_impact", "schedule_impact"}
        )
        compact.setdefault("impacts", impacts)
    return compact, tuple(warnings)


def load_document_bytes(data: bytes, filename: str) -> Document:
    """Load supported document bytes into a typed normalized document."""
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError(
            f"{extension or 'This file type'} is unsupported. Use .md, .txt, or .docx."
        )
    raw = (
        _read_docx(data, filename)
        if extension == ".docx"
        else _decode_text(data, filename)
    )
    text = normalize_text(raw)
    if not text:
        raise DocumentParseError(f"{filename} contains no readable text.")
    sections, warnings = segment_sections(text)
    content_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()
    number = re.search(r"\bCO-\d{3}\b", text, flags=re.IGNORECASE)
    document_id = number.group(0).upper() if number else Path(filename).stem
    return Document(
        document_id=document_id,
        filename=Path(filename).name,
        path=None,
        text=text,
        sections=sections,
        warnings=warnings,
        content_hash=content_hash,
    )


def load_document(path: Path) -> Document:
    """Load a document path without writing or mutating the source file."""
    try:
        document = load_document_bytes(path.read_bytes(), path.name)
    except OSError as exc:
        raise DocumentParseError(f"Could not read {path.name}: {exc}") from exc
    return Document(
        document_id=document.document_id,
        filename=document.filename,
        path=path,
        text=document.text,
        sections=document.sections,
        warnings=document.warnings,
        content_hash=document.content_hash,
    )


def load_documents(paths: list[Path]) -> tuple[Document, ...]:
    """Load multiple documents deterministically in filename order."""
    return tuple(load_document(path) for path in sorted(paths, key=lambda p: p.name))
