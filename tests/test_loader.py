"""Document-loader robustness tests."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as WordDocument

from co_copilot.document_loader import (
    load_document,
    load_document_bytes,
    normalize_text,
    segment_sections,
)
from co_copilot.exceptions import DocumentParseError


def test_loads_demo_markdown_with_sections(project_root: Path) -> None:
    document = load_document(project_root / "data/change_orders/CO-001.md")
    assert document.document_id == "CO-001"
    assert {"header", "description", "basis", "impacts"} <= set(document.sections)
    assert len(document.content_hash) == 64


def test_messy_unknown_structure_degrades_to_full_text() -> None:
    text = "\ufeff  Change   order CO-999\r\nUnstructured but readable narrative. " * 4
    document = load_document_bytes(text.encode(), "messy.txt")
    assert document.sections == {"full_text": document.text}
    assert document.warnings
    assert "\r" not in document.text


def test_docx_ingestion_uses_memory_bytes() -> None:
    word = WordDocument()
    word.add_heading("Change Order CO-777", level=1)
    word.add_paragraph("HEADER")
    word.add_paragraph("Contractor: Synthetic Contractor")
    word.add_paragraph("EVENT DESCRIPTION")
    word.add_paragraph("Additional pipework required by instruction.")
    buffer = BytesIO()
    word.save(buffer)
    loaded = load_document_bytes(buffer.getvalue(), "sample.docx")
    assert loaded.document_id == "CO-777"
    assert "Synthetic Contractor" in loaded.text


@pytest.mark.parametrize("name", ["bad.pdf", "archive.zip", "no_extension"])
def test_rejects_unsupported_extension(name: str) -> None:
    with pytest.raises(DocumentParseError, match="unsupported"):
        load_document_bytes(b"not a supported document", name)


def test_rejects_empty_and_invalid_docx() -> None:
    with pytest.raises(DocumentParseError, match="no readable text"):
        load_document_bytes(b"", "empty.txt")
    with pytest.raises(DocumentParseError, match="not a readable Word"):
        load_document_bytes(b"not a zip", "broken.docx")


def test_helpers_normalize_and_segment() -> None:
    assert normalize_text("A   B\r\nC") == "A B\nC"
    sections, warnings = segment_sections(
        "HEADER\nDate: 2026-01-01\nCOST IMPACT\nUSD 2"
    )
    assert not warnings
    assert sections["header"].startswith("Date")
    assert "impacts" in sections
